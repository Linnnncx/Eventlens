from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "out"
OUTPUT = OUT_DIR / "EventLens_项目介绍与使用指南.docx"
WORKBENCH_IMAGE = OUT_DIR / "eventlens-workbench.png"
MOBILE_IMAGE = OUT_DIR / "eventlens-mobile.png"

NAVY = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GOLD = "B28A38"
INK = "1C2530"
GRAY = "66717E"
LIGHT = "F4F6F9"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF8E8"
GREEN = "237A46"
RED = "9B2C2C"
WHITE = "FFFFFF"


def set_run_font(run, size=None, color=INK, bold=None, italic=None, latin="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = latin
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_cell_border(cell, color="D5DBE3", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    tr_pr.append(hdr)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()._r
    r.extend([fld_begin, instr, fld_sep, text, fld_end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9, GRAY)


def add_para(doc, text="", *, style=None, size=None, color=INK, bold=False, italic=False,
             align=None, before=0, after=8, line=1.333, keep_with_next=False):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r, size, color, bold, italic)
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next
    if align is not None:
        p.alignment = align
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    r = p.add_run(text)
    set_run_font(r, 10.5, INK)
    return p


def add_numbered(doc, title, detail):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.208
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    r1 = p.add_run(title + "：")
    set_run_font(r1, 10.5, DARK_BLUE, True)
    r2 = p.add_run(detail)
    set_run_font(r2, 10.5, INK)
    return p


def add_callout(doc, label, text, fill=LIGHT, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), accent)
    left.set(qn("w:space"), "6")
    borders.append(left)
    p_pr.append(borders)
    r1 = p.add_run(label + "  ")
    set_run_font(r1, 10.5, accent, True)
    r2 = p.add_run(text)
    set_run_font(r2, 10.5, INK)
    return p


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "10141B")
    p_pr.append(shd)
    for idx, line in enumerate(lines):
        r = p.add_run(line)
        set_run_font(r, 9, "E6EDF3", latin="Consolas", east_asia="Microsoft YaHei")
        if idx < len(lines) - 1:
            r.add_break()
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        set_run_font(run, {1: 16, 2: 13, 3: 12}[level], {1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], True)
    return p


def add_figure(doc, image_path, width, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    picture_run = p.add_run()
    picture_run.add_picture(str(image_path), width=Inches(width))
    doc_pr = picture_run._r.xpath(".//wp:docPr")[0]
    doc_pr.set("descr", caption)
    doc_pr.set("title", caption.split("｜", 1)[0])
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_before = Pt(0)
    c.paragraph_format.space_after = Pt(10)
    r = c.add_run(caption)
    set_run_font(r, 9, GRAY, italic=True)


def add_comparison_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, PALE_BLUE)
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, 9.5, DARK_BLUE, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if idx == 0:
                r = p.add_run(text)
                set_run_font(r, 9.5, DARK_BLUE, True)
            else:
                r = p.add_run(text)
                set_run_font(r, 9.2, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def configure_section(section, first=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = first


def set_running_furniture(section, title="EventLens 项目介绍与使用指南"):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    set_run_font(r, 9, GRAY, True)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    configure_section(section, first=True)

    # Cover: narrative_proposal preset with editorial_cover title override.
    add_para(doc, "PRODUCT BRIEF · 2026", size=10, color=GOLD, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=88, after=18, line=1.0)
    add_para(doc, "EventLens", size=31, color=NAVY, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=4, line=1.0)
    add_para(doc, "项目介绍与使用指南", size=22, color=DARK_BLUE, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=18, line=1.0)
    add_para(doc, "事件驱动型美股研究与模拟交易工作台", size=14, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=70, line=1.1)
    add_callout(doc, "一句话定位", "把新闻事件、价格行为、技术指标、AI 解释与模拟交易放进同一条决策链路，让用户从“看到消息”自然走到“验证影响、评估风险、复盘结果”。", fill=PALE_GOLD, accent=GOLD)
    add_para(doc, "基于当前项目代码与本地界面梳理", size=9.5, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=36, after=3, line=1.0)
    add_para(doc, "版本日期：2026 年 8 月 4 日", size=9.5, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0)

    doc.add_page_break()
    set_running_furniture(section)

    add_heading(doc, "执行摘要", 1)
    add_para(doc, "EventLens 面向希望理解“事件如何影响价格”的用户。它不是单纯的行情看板，也不是把聊天机器人嵌入券商页面，而是围绕事件研究构建的完整工作流：市场扫描 → 标的研究 → 事件与 K 线对齐 → 技术面与新闻面联合分析 → 下单前风险预览 → 模拟成交 → 资产与交易复盘。")
    add_callout(doc, "核心价值", "把分散在行情软件、新闻聚合、技术分析、AI 问答和模拟交易工具中的动作，收拢到一个可解释、可验证、可复盘的研究环境中。")

    add_heading(doc, "阅读导航", 2)
    add_comparison_table(
        doc,
        ["模块", "回答的问题", "主要读者"],
        [
            ("产品介绍", "EventLens 是什么、解决什么问题", "潜在用户、合作方、评审者"),
            ("使用方法", "如何启动、研究标的、分析事件并模拟交易", "首次使用者、演示人员"),
            ("目标群体", "哪些用户最能获得价值", "产品与市场团队"),
            ("差异化与创新", "为什么不是另一个行情或 AI 工具", "产品经理、投资人、技术评审"),
            ("架构与边界", "系统如何实现、当前适用范围是什么", "开发者、部署与运维人员"),
        ],
        [1700, 4300, 3360],
    )

    add_heading(doc, "1. 项目概览", 1)
    add_heading(doc, "1.1 产品要解决的问题", 2)
    add_para(doc, "传统研究流程往往被切割成多个孤立页面：用户先在新闻平台看到事件，再切到行情软件寻找对应时间点，手工判断价格与量能变化，随后到另一处计算仓位、风险与交易成本。信息虽然丰富，但缺少统一的时间坐标和行动上下文。")
    add_para(doc, "EventLens 以“事件”为入口、以“时间对齐”为骨架、以“模拟交易”为闭环，将新闻、K 线、技术指标、事件反应与风险预览组织到同一工作台中。产品强调研究与验证，不承诺预测股价，也不输出确定性的买卖建议。")

    add_heading(doc, "1.2 产品定位", 2)
    add_comparison_table(
        doc,
        ["维度", "EventLens 的定位"],
        [
            ("产品类别", "事件驱动型美股研究与模拟交易工作台"),
            ("核心对象", "新闻事件及其前后价格、成交量与技术结构变化"),
            ("核心行为", "观察、对齐、分析、预览、模拟、复盘"),
            ("数据范围", "美股行情、K 线、新闻、事件标签与模拟账户数据"),
            ("安全边界", "研究与产品演示用途；不构成投资建议，不接入真实券商成交"),
        ],
        [1900, 7460],
    )

    add_heading(doc, "1.3 产品闭环", 2)
    for title, detail in [
        ("扫描市场", "通过指数、今日涨跌热点、核心异动与自选列表快速发现关注标的"),
        ("进入工作台", "在同一视图中查看价格、K 线、技术指标、新闻锚点和市场状态"),
        ("理解事件", "查看事件方向、重要性、类型、发布时间及事件前后反应窗口"),
        ("形成解释", "以规则引擎或可选 LLM 联合分析技术面与新闻面，保留事实依据"),
        ("预演交易", "在提交模拟订单前计算金额、费用、现金、仓位与集中度风险"),
        ("复盘结果", "通过持仓、委托、成交、资产曲线与已平仓排行回看决策结果"),
    ]:
        add_numbered(doc, title, detail)

    add_heading(doc, "2. 功能全景", 1)
    add_heading(doc, "2.1 市场首页：从全局异动进入研究", 2)
    add_bullet(doc, "展示总资产、现金、市场开闭市状态、持仓数量与自选数量。")
    add_bullet(doc, "聚合道琼斯、标普 500、纳斯达克等指数的实时/准实时表现。")
    add_bullet(doc, "以涨跌热力图突出当日强弱标的，并提供 Core Movers、自选、持仓与近期订单入口。")
    add_bullet(doc, "支持股票代码/公司搜索，缩短从发现到研究的路径。")

    add_heading(doc, "2.2 事件研究工作台：产品的核心界面", 2)
    if WORKBENCH_IMAGE.exists():
        add_figure(doc, WORKBENCH_IMAGE, 6.5, "图 1｜桌面工作台：自选与市场列表、事件锚点 K 线、指标区、新闻/交易区在同一屏协同")
    add_bullet(doc, "多周期 K 线：1Min、5Min、15Min、1Hour、4Hour、1Day、1Month。")
    add_bullet(doc, "事件锚点：将新闻按市场时区和周期自动映射到对应 K 线桶，颜色区分利好、利空与中性。")
    add_bullet(doc, "事件反应：计算事件前 5/30 分钟、后 5/30/60 分钟涨跌，以及最大上行、最大回撤和量能倍数。")
    add_bullet(doc, "技术工具：均线、EMA、BOLL、VWAP、SAR、支撑/压力、Donchian，以及 MACD、RSI、KDJ、CCI、ATR、OBV、ADX、MFI 等副图指标。")
    add_bullet(doc, "图表交互：标价、水平线、趋势线、射线、垂直线、矩形和斐波那契等画线工具；布局可拖拽调整并一键重置。")

    add_heading(doc, "2.3 AI 区间分析与新闻解释", 2)
    add_para(doc, "用户选择标的、周期和时间区间后，系统将该窗口内的价格结构、技术指标与新闻事件联合整理成分析上下文。默认规则引擎可离线工作，也可配置 OpenAI、DeepSeek 或 Qwen 的 OpenAI 兼容接口生成更自然的解释。")
    add_callout(doc, "可解释性原则", "LLM 负责解释已计算的事实，不替代价格计算、事件反应计算或硬性风险规则；云模型不可用时仍可降级到规则模板。")

    add_heading(doc, "2.4 模拟交易与风险预览", 2)
    add_bullet(doc, "支持买入/卖出、市场价/限价及扩展订单类型的交互入口。")
    add_bullet(doc, "下单前展示预计金额、预计手续费、成交价格、数量、下单后现金、单一持仓占比与订单占总资产比例。")
    add_bullet(doc, "基于单一持仓集中度、行业集中度、现金比例、订单规模、事件后波动、量能异常、止损设置和重大事件新鲜度生成风险提示。")
    add_bullet(doc, "模拟成交后自动更新现金、持仓成本、订单、成交记录、资产曲线与已实现盈亏。")

    add_heading(doc, "2.5 移动端：重组后的移动体验", 2)
    if MOBILE_IMAGE.exists():
        add_figure(doc, MOBILE_IMAGE, 2.55, "图 2｜移动端行情页：指数、资产、自选与底部四栏导航针对小屏重新组织")
    add_para(doc, "移动端并非简单压缩桌面布局，而是独立采用“行情 / 资讯 / 交易 / 我的”四栏导航。股票详情保留 K 线、新闻锚点、指标设置和快捷下单；交易页集中展示资产、持仓与委托；“我的”提供资产曲线、已平仓排行、数据源与模型状态。")

    add_heading(doc, "3. 使用方法", 1)
    add_heading(doc, "3.1 本地启动", 2)
    add_para(doc, "环境要求：Python 3.11+、Node.js 20+。首次使用先复制环境变量示例，再分别启动后端和前端。")
    add_code_block(doc, [
        "# 1. 在 eventlens 根目录创建配置",
        "cp .env.example .env",
        "",
        "# 2. 后端（新终端）",
        "cd backend",
        "python -m venv .venv",
        ".\\.venv\\Scripts\\activate   # Windows",
        "pip install -r requirements.txt",
        "uvicorn app.main:app --reload --port 8000",
        "",
        "# 3. 前端（新终端）",
        "cd frontend",
        "npm install",
        "npm run dev",
    ])
    add_para(doc, "启动后访问前端 http://localhost:5173；后端 API 为 http://localhost:8000；Swagger 文档为 http://localhost:8000/docs。窄屏会自动切换到移动端外壳。", size=10.5)

    add_heading(doc, "3.2 首次配置建议", 2)
    add_comparison_table(
        doc,
        ["配置项", "建议", "作用"],
        [
            ("行情 Provider", "yfinance；有 Alpaca Key 时可切换", "报价、K 线、搜索与实时刷新"),
            ("新闻 Provider", "merged", "并发合并 Google News、Finnhub、Yahoo 并去重"),
            ("Fixture 模式", "演示或离线时开启", "在外部数据不可用时稳定展示完整流程"),
            ("LLM Provider", "先用 rules，再按需配置云模型", "保证无 Key 也可运行，配置后增强自然语言解释"),
            ("初始资金", "按演示场景设置 INITIAL_CASH", "决定模拟账户基准"),
        ],
        [1750, 3000, 4610],
    )

    add_heading(doc, "3.3 推荐的桌面端研究流程", 2)
    for title, detail in [
        ("选择标的", "从热点、自选、持仓或搜索进入工作台"),
        ("选择周期", "事件密集的盘中研究优先 5Min/15Min；趋势研究可使用 1Hour/1Day"),
        ("观察锚点", "悬停或点击带事件标记的 K 线，查看对应新闻及事件反应"),
        ("交叉验证", "同时观察均线、RSI、MACD、成交量等技术结构，避免只看标题情绪"),
        ("运行区间分析", "设定开始/结束日期，生成技术面与新闻面的联合报告"),
        ("预览模拟订单", "输入方向、数量或金额，检查费用、现金变化、仓位权重和风险提示"),
        ("提交与复盘", "在持仓、委托、成交及资产曲线中追踪结果"),
    ]:
        add_numbered(doc, title, detail)

    add_heading(doc, "3.4 Fixture 演示流程", 2)
    add_callout(doc, "最快演示路径", "将 FIXTURE_MODE=true，打开 NVDA 工作台并选择 5Min；点击有新闻锚点的 K 线查看事件反应，再提交一笔模拟订单，观察持仓、成本线和资产数据更新。", fill=PALE_GOLD, accent=GOLD)

    add_heading(doc, "3.5 生产部署", 2)
    add_para(doc, "仓库已提供 Docker Compose、前后端 Dockerfile 与 Nginx 反向代理，可部署到单机 VPS。生产环境建议使用 docker-compose.prod.yml 与独立的 .env.production，并只对公网开放 80/443。后端 8000 端口应保留在容器网络内。")
    add_code_block(doc, [
        "cp .env.production.example .env.production",
        "docker compose -f docker-compose.prod.yml --env-file .env.production up -d --build",
        "docker compose -f docker-compose.prod.yml ps",
    ])

    add_heading(doc, "4. 目标群体与典型场景", 1)
    add_comparison_table(
        doc,
        ["目标群体", "主要诉求", "典型使用场景"],
        [
            ("事件驱动型交易研究者", "快速确认新闻发生后价格与量能是否真实响应", "财报、产品发布、监管、并购、宏观消息后的反应研究"),
            ("主动投资者与自选股用户", "减少新闻与图表来回切换，建立结构化观察流程", "盘前准备、盘中跟踪、盘后复盘"),
            ("量化/金融科技产品学习者", "理解事件对齐、指标计算、风险规则与数据降级", "课程项目、作品集、原型验证、技术演示"),
            ("交易新手与教育场景", "在不承担真实资金风险的情况下学习订单与仓位", "模拟下单、手续费理解、集中度与止损教育"),
            ("产品评审与合作方", "快速评估事件研究产品的完整性与可扩展性", "产品路演、合作沟通、内部评审"),
        ],
        [2050, 3500, 3810],
    )
    add_callout(doc, "不适用人群", "需要真实券商成交、超低延迟专业盘口、机构级合规审计或自动化实盘策略执行的用户，不应把当前版本视为生产级交易终端。", fill="FDECEC", accent=RED)

    add_heading(doc, "5. 设计差异化", 1)
    add_heading(doc, "5.1 从“看行情”转向“理解事件—验证反应”", 2)
    add_para(doc, "多数行情产品以股票为中心，新闻只是旁栏信息；EventLens 将事件变成 K 线的一等对象。每条新闻被映射到具体时间桶，并与价格变化、成交量和技术结构形成可追踪关系。用户看到的不只是“发生了什么”，还包括“市场是否、何时、以多大幅度作出反应”。")

    add_heading(doc, "5.2 把研究与行动放在同一个上下文", 2)
    add_para(doc, "工作台不要求用户在新闻、图表、订单和账户页面之间反复切换。选中的事件、当前标的、图表周期、标记价格与模拟订单共享上下文，使“研究—判断—预演—复盘”连续发生。")

    add_heading(doc, "5.3 AI 不是结论机器，而是事实解释层", 2)
    add_para(doc, "系统先由代码计算行情、指标、事件反应与风险事实，再让规则模板或 LLM 进行归纳。这样的分工降低了模型编造数字、遮蔽不确定性或在网络异常时让核心流程失效的风险。")

    add_heading(doc, "5.4 移动端不是桌面端的缩小版", 2)
    add_para(doc, "桌面端强调多面板并行观察，移动端强调任务切换和单手操作。独立路由、底部导航、下拉刷新、底部 Sheet、触觉反馈、可拖拽自选与全屏图表体现了针对设备场景的重新设计。")

    add_heading(doc, "5.5 对比视角", 2)
    add_comparison_table(
        doc,
        ["对比维度", "常见行情/新闻工具", "EventLens"],
        [
            ("信息组织", "行情、新闻、交易相互分离", "围绕同一标的与同一时间窗口联动"),
            ("新闻与图表", "新闻列表与 K 线弱关联", "自适应周期对齐并形成可点击锚点"),
            ("事件影响", "依赖用户主观观察", "提供标准化前后反应窗口与量能倍数"),
            ("AI 使用", "直接生成观点或摘要", "先计算事实，再解释；可规则降级"),
            ("交易安全", "直接下单或只提供简单模拟", "下单前集中度、现金与事件波动预览"),
            ("跨端设计", "响应式压缩桌面页面", "桌面并行工作台 + 移动任务式导航"),
            ("数据可靠性", "单一数据源失败即中断", "Provider、缓存、Fixture 与结构化错误多级降级"),
        ],
        [1750, 3500, 4110],
    )

    add_heading(doc, "6. 创新点", 1)
    innovation_items = [
        ("周期自适应的事件—K 线对齐", "新闻时间先转换到市场时区，再按 1Min 至 1Month 的周期规则映射到 K 线桶；日级还可聚合多条事件，解决跨周期观察时锚点漂移的问题。"),
        ("标准化事件反应窗口", "以事件时刻为中心，自动计算前后涨跌、最大上行、最大回撤、当日高低点与量能倍数，把主观“好像涨了”转化为可比较事实。"),
        ("技术面 × 新闻面的区间级联合分析", "不是只分析单条新闻，而是把指定时间区间内的行情结构、指标状态和事件集合共同纳入解释。"),
        ("风险预览前置", "在模拟订单确认前就量化仓位集中度、现金比例、订单规模及事件后波动，让风险提示成为交易流程的一部分，而非成交后的报告。"),
        ("可替换 AI 与可持续降级", "同一分析接口支持规则模板、OpenAI、DeepSeek 与 Qwen；无 Key、超时或外部服务不可用时，核心功能仍可继续。"),
        ("共享行情单航班缓存", "REST 报价、订单簿与实时广播共享按标的缓存和锁，同一时刻只触发一次上游请求，降低重复请求与首屏等待。"),
        ("事件弹幕与价格上下文", "工作台提供与标的共同存在的轻量弹幕通道，为未来的事件讨论、观点标注和协作研究预留交互层。"),
    ]
    for title, detail in innovation_items:
        add_heading(doc, title, 2)
        add_para(doc, detail, size=10.5, after=6)

    add_heading(doc, "7. 技术架构与工程特点", 1)
    add_comparison_table(
        doc,
        ["层级", "技术与职责"],
        [
            ("Web 前端", "React 18 + TypeScript + Vite；TanStack Query 管理服务端状态，Zustand 管理工作台与账户状态，Lightweight Charts 承载图表"),
            ("移动端", "与桌面端共享 API 与核心数据结构，但采用独立页面、组件与交互组织"),
            ("API 后端", "FastAPI + Pydantic；提供行情、新闻、事件、AI 分析、账户、订单、成交、弹幕与配置接口"),
            ("数据与持久化", "SQLAlchemy + SQLite；保存账户、持仓、订单、成交、新闻缓存、运行时配置与资产曲线"),
            ("Provider 层", "行情支持 yfinance / Alpaca / Fixture；新闻支持 merged / Finnhub / Yahoo / Google News / Fixture"),
            ("AI 层", "规则模板或 OpenAI 兼容接口；支持 OpenAI、DeepSeek、Qwen"),
            ("部署层", "Docker Compose + Nginx；前端同域反代 /api 与 /ws"),
        ],
        [1900, 7460],
    )

    add_heading(doc, "7.1 关键工程取舍", 2)
    add_bullet(doc, "Provider 抽象让外部数据源可替换，并为测试与离线演示保留 Fixture。")
    add_bullet(doc, "缓存与持久化优先返回可用数据，再后台刷新，改善外部数据源不稳定时的体验。")
    add_bullet(doc, "订单预览使用确定性规则快速计算，不让云模型延迟阻塞确认按钮。")
    add_bullet(doc, "API Key 仅保存在后端，前端不会读取或明文回传。")
    add_bullet(doc, "前后端均包含自动化测试，覆盖事件对齐、指标、移动图表点击、行情与新闻调度、订单预览等关键路径。")

    add_heading(doc, "8. 产品边界、风险与已知限制", 1)
    add_comparison_table(
        doc,
        ["边界/限制", "说明与建议"],
        [
            ("非真实交易", "所有订单均为模拟成交，不连接真实券商；不可作为实盘成交记录"),
            ("数据时效", "行情和新闻可能延迟、来自缓存或 Fixture；重要决策应回查权威来源"),
            ("数据源限制", "yfinance 日内历史范围有限，WebSocket 不稳定时会降级为轮询"),
            ("存储扩展性", "当前 SQLite 适合演示与小流量单机部署；多实例和高并发应迁移到 Postgres 等共享数据库"),
            ("外部网络", "真实行情与新闻依赖 Yahoo、Finnhub 等外部服务，部署环境需具备稳定出网能力"),
            ("合规属性", "AI 分析、事件标签和风险提示均不构成投资建议，也不保证收益"),
        ],
        [2100, 7260],
    )
    add_callout(doc, "风险声明", "EventLens 仅供研究、学习与产品演示。市场数据可能延迟或不准确；任何分析、分类和风险说明都不应被理解为买卖建议或收益承诺。", fill="FDECEC", accent=RED)

    add_heading(doc, "9. 产品价值总结", 1)
    add_para(doc, "EventLens 的价值不在于堆叠更多行情指标，而在于建立一条清晰的事件研究链路：让时间成为新闻与市场反应的共同坐标，让代码负责可验证的计算，让 AI 负责解释，让风险检查发生在行动之前，让模拟账户承接复盘。")
    add_para(doc, "这种设计使它同时具备三种属性：对研究者，它是事件验证工具；对学习者，它是低风险交易实验环境；对产品与技术评审者，它是一个从数据接入、交互设计到降级与部署都较完整的金融科技原型。")
    add_callout(doc, "最终定位", "一个以事件为核心、以可解释性为原则、以模拟闭环为边界的跨端美股研究工作台。", fill=PALE_GOLD, accent=GOLD)

    add_heading(doc, "附录 A｜常用环境变量", 1)
    add_comparison_table(
        doc,
        ["变量", "用途"],
        [
            ("FRONTEND_ORIGIN", "前端源站与 CORS 配置"),
            ("MARKET_DATA_PROVIDER", "yfinance / alpaca / fixture"),
            ("NEWS_PROVIDER", "merged / finnhub / yfinance / alpaca / fixture"),
            ("REALTIME_PROVIDER", "yfinance / alpaca / fixture"),
            ("FIXTURE_MODE", "强制使用演示数据"),
            ("LLM_PROVIDER", "rules 或云模型 Provider"),
            ("FINNHUB_API_KEY", "启用 Finnhub 新闻"),
            ("ALPACA_API_KEY / SECRET", "启用 Alpaca 数据源"),
            ("INITIAL_CASH", "模拟账户初始资金"),
        ],
        [3100, 6260],
    )

    # Core properties and output.
    doc.core_properties.title = "EventLens 项目介绍与使用指南"
    doc.core_properties.subject = "事件驱动型美股研究与模拟交易工作台"
    doc.core_properties.author = "EventLens Project"
    doc.core_properties.keywords = "EventLens, 事件驱动, 美股研究, 模拟交易, AI 分析"
    doc.core_properties.comments = "基于 2026-08-04 项目代码与本地界面生成"

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
